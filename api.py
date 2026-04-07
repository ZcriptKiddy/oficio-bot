from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.shared import Pt
import uuid
import os

app = Flask(__name__)

# 🔤 FUNCIÓN: aplicar formato Arial Narrow 10
def aplicar_formato(run):
    run.font.name = "Arial Narrow"
    run.font.size = Pt(10)


# 🔁 REEMPLAZO EN TEXTO (párrafos, tablas y encabezados)
def reemplazar(doc, data):

    # 📄 PÁRRAFOS
    for p in doc.paragraphs:
        for key, value in data.items():
            if f"{{{{{key}}}}}" in p.text:
                p.text = p.text.replace(f"{{{{{key}}}}}", str(value))

        # 🔤 aplicar formato a cada run
        for run in p.runs:
            aplicar_formato(run)

    # 📊 TABLAS
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in data.items():
                        if f"{{{{{key}}}}}" in p.text:
                            p.text = p.text.replace(f"{{{{{key}}}}}", str(value))

                    for run in p.runs:
                        aplicar_formato(run)

    # 📑 ENCABEZADOS
    for section in doc.sections:
        header = section.header
        for p in header.paragraphs:
            for key, value in data.items():
                if f"{{{{{key}}}}}" in p.text:
                    p.text = p.text.replace(f"{{{{{key}}}}}", str(value))

            for run in p.runs:
                aplicar_formato(run)


# 🌐 Endpoint base
@app.route("/")
def home():
    return "API activa"


# 📄 GENERAR DOCUMENTO
@app.route("/generar", methods=["POST"])
def generar():
    try:
        data = request.json

        if not data:
            return jsonify({"error": "No se recibieron datos"}), 400

        doc = Document("plantilla.docx")

        reemplazar(doc, data)

        nombre_archivo = f"oficio_{uuid.uuid4()}.docx"
        doc.save(nombre_archivo)

        return send_file(nombre_archivo, as_attachment=True)

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# 🚀 LOCAL
if __name__ == "__main__":
    app.run(host="127.0.0.1", port=5000, debug=True)